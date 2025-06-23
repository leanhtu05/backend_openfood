from fastapi import APIRouter, Request, Depends, HTTPException, Query, Form
from fastapi.responses import HTMLResponse, RedirectResponse, StreamingResponse, Response
from fastapi.templating import Jinja2Templates
from typing import List, Dict, Optional, Any
import os
from datetime import datetime, timedelta
import json
import csv
import io
import secrets
import time
try:
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment
    from openpyxl.utils.dataframe import dataframe_to_rows
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    WORD_AVAILABLE = True
except ImportError:
    WORD_AVAILABLE = False

# Temporary tokens for download (in-memory storage)
download_tokens = {}

# Import services
from services.firestore_service import firestore_service
from middleware.auth import (
    authenticate_admin,
    create_admin_session,
    get_current_admin,
    delete_admin_session,
    require_admin_auth
)
from auth_utils import get_current_user, TokenPayload

router = APIRouter(prefix="/admin", tags=["Admin"])

# ==================== AUTHENTICATION ROUTES ====================

@router.get("/login", response_class=HTMLResponse)
async def admin_login_page(request: Request, error: str = None, success: str = None):
    """Hi·ªÉn th·ªã trang ƒëƒÉng nh·∫≠p admin"""
    # N·∫øu ƒë√£ ƒëƒÉng nh·∫≠p r·ªìi th√¨ redirect v·ªÅ dashboard
    if get_current_admin(request):
        return RedirectResponse(url="/admin/", status_code=302)

    templates = get_templates()
    return templates.TemplateResponse("admin/login.html", {
        "request": request,
        "error": error,
        "success": success
    })

@router.get("/fast-login", response_class=HTMLResponse)
async def admin_fast_login_page(
    request: Request,
    error: str = None,
    success: str = None,
    templates: Jinja2Templates = Depends(get_templates)
):
    """üöÄ Trang ƒëƒÉng nh·∫≠p admin t·ªëi ∆∞u h√≥a"""
    # N·∫øu ƒë√£ ƒëƒÉng nh·∫≠p r·ªìi th√¨ redirect v·ªÅ fast dashboard
    if get_current_admin(request):
        return RedirectResponse(url="/admin/fast-dashboard", status_code=302)

    return templates.TemplateResponse("admin/fast_login.html", {
        "request": request,
        "error": error,
        "success": success
    })

@router.post("/login")
async def admin_login(request: Request, username: str = Form(...), password: str = Form(...)):
    """X·ª≠ l√Ω ƒëƒÉng nh·∫≠p admin"""
    try:
        print(f"[AUTH] Admin login attempt: {username}")

        if authenticate_admin(username, password):
            # T·∫°o session
            session_token = create_admin_session(username)
            print(f"[AUTH] Admin login successful: {username}")

            # üöÄ Redirect v·ªÅ fast dashboard v·ªõi session cookie
            response = RedirectResponse(url="/admin/fast-dashboard", status_code=302)
            response.set_cookie(
                key="admin_session",
                value=session_token,
                max_age=86400,  # 24 hours
                httponly=True,
                secure=False,  # Set to True in production with HTTPS
                samesite="lax"
            )
            return response
        else:
            print(f"[AUTH] Admin login failed: {username}")
            return RedirectResponse(
                url="/admin/fast-login?error=T√™n ƒëƒÉng nh·∫≠p ho·∫∑c m·∫≠t kh·∫©u kh√¥ng ƒë√∫ng",
                status_code=302
            )
    except Exception as e:
        print(f"[AUTH] Admin login error: {str(e)}")
        return RedirectResponse(
            url="/admin/fast-login?error=C√≥ l·ªói x·∫£y ra khi ƒëƒÉng nh·∫≠p",
            status_code=302
        )

@router.post("/fast-login")
async def admin_fast_login(request: Request, username: str = Form(...), password: str = Form(...)):
    """üöÄ X·ª≠ l√Ω ƒëƒÉng nh·∫≠p admin t·ªëi ∆∞u h√≥a"""
    try:
        print(f"[FAST-AUTH] Admin login attempt: {username}")

        if authenticate_admin(username, password):
            # T·∫°o session
            session_token = create_admin_session(username)
            print(f"[FAST-AUTH] Admin login successful: {username}")

            # Redirect v·ªÅ fast dashboard v·ªõi session cookie
            response = RedirectResponse(url="/admin/fast-dashboard", status_code=302)
            response.set_cookie(
                key="admin_session",
                value=session_token,
                max_age=86400,  # 24 hours
                httponly=True,
                secure=False,  # Set to True in production with HTTPS
                samesite="lax"
            )
            return response
        else:
            print(f"[FAST-AUTH] Admin login failed: {username}")
            return RedirectResponse(
                url="/admin/fast-login?error=T√™n ƒëƒÉng nh·∫≠p ho·∫∑c m·∫≠t kh·∫©u kh√¥ng ƒë√∫ng",
                status_code=302
            )
    except Exception as e:
        print(f"[FAST-AUTH] Admin login error: {str(e)}")
        return RedirectResponse(
            url="/admin/fast-login?error=C√≥ l·ªói x·∫£y ra khi ƒëƒÉng nh·∫≠p",
            status_code=302
        )

@router.get("/logout")
async def admin_logout(request: Request):
    """ƒêƒÉng xu·∫•t admin"""
    session_token = request.cookies.get("admin_session")
    if session_token:
        delete_admin_session(session_token)

    response = RedirectResponse(url="/admin/login?success=ƒê√£ ƒëƒÉng xu·∫•t th√†nh c√¥ng", status_code=302)
    response.delete_cookie("admin_session")
    return response

# ==================== ADMIN PAGES ====================

# Template instance
def get_templates():
    base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    return Jinja2Templates(directory=os.path.join(base_dir, "templates"))

# Import food_items from openfood_router (fallback)
try:
    from routers.openfood_router import food_items as fallback_food_items
except ImportError:
    fallback_food_items = []

def get_foods_data():
    """L·∫•y d·ªØ li·ªáu foods t·ª´ Firebase, fallback v·ªÅ d·ªØ li·ªáu m·∫´u n·∫øu c·∫ßn"""
    try:
        # Th·ª≠ l·∫•y t·ª´ Firebase tr∆∞·ªõc
        foods = firestore_service.get_all_foods()
        if foods:
            return foods
        else:
            # Fallback v·ªÅ d·ªØ li·ªáu m·∫´u
            return fallback_food_items
    except Exception as e:
        print(f"Error getting foods from Firebase: {str(e)}")
        # Fallback v·ªÅ d·ªØ li·ªáu m·∫´u
        return fallback_food_items

def get_system_stats():
    """L·∫•y th·ªëng k√™ t·ªïng quan c·ªßa h·ªá th·ªëng"""
    try:
        # Th·ªëng k√™ m√≥n ƒÉn t·ª´ Firebase
        foods_data = get_foods_data()
        total_foods = len(foods_data)
        
        # Th·ªëng k√™ ng∆∞·ªùi d√πng (t·ª´ Firestore)
        try:
            users = firestore_service.get_all_users()
            active_users = len([u for u in users if u.get('last_login')])
        except:
            active_users = 0
        
        # Th·ªëng k√™ meal plans
        try:
            meal_plans = firestore_service.get_all_meal_plans()
            total_meal_plans = len(meal_plans)
        except:
            total_meal_plans = 0
        
        # API calls h√¥m nay (gi·∫£ l·∫≠p)
        api_calls_today = 150  # C√≥ th·ªÉ implement tracking th·ª±c t·∫ø
        
        return {
            "total_foods": total_foods,
            "active_users": active_users,
            "total_meal_plans": total_meal_plans,
            "api_calls_today": api_calls_today
        }
    except Exception as e:
        print(f"Error getting system stats: {str(e)}")
        return {
            "total_foods": 0,
            "active_users": 0,
            "total_meal_plans": 0,
            "api_calls_today": 0
        }

def get_recent_activities():
    """L·∫•y ho·∫°t ƒë·ªông g·∫ßn ƒë√¢y t·ª´ Firebase"""
    try:
        activities = []

        # L·∫•y meal plans g·∫ßn ƒë√¢y
        recent_meal_plans = firestore_service.get_all_meal_plans()
        for plan in recent_meal_plans[:3]:  # L·∫•y 3 meal plans g·∫ßn nh·∫•t
            activities.append({
                "action": "T·∫°o meal plan",
                "description": f"K·∫ø ho·∫°ch b·ªØa ƒÉn cho user {plan.get('user_id', 'Unknown')[:8]}...",
                "timestamp": plan.get('created_at', 'Kh√¥ng r√µ')
            })

        # L·∫•y ng∆∞·ªùi d√πng m·ªõi
        recent_users = firestore_service.get_all_users()
        for user in recent_users[:2]:  # L·∫•y 2 users g·∫ßn nh·∫•t
            activities.append({
                "action": "Ng∆∞·ªùi d√πng ƒëƒÉng k√Ω",
                "description": f"Ng∆∞·ªùi d√πng m·ªõi: {user.get('email', 'Unknown')}",
                "timestamp": user.get('created_at', 'Kh√¥ng r√µ')
            })

        # S·∫Øp x·∫øp theo th·ªùi gian (n·∫øu c√≥)
        return activities[:5]  # Tr·∫£ v·ªÅ 5 ho·∫°t ƒë·ªông g·∫ßn nh·∫•t

    except Exception as e:
        print(f"Error getting recent activities: {str(e)}")
        # Fallback to mock data
        return [
            {
                "action": "H·ªá th·ªëng kh·ªüi ƒë·ªông",
                "description": "Server ƒë√£ kh·ªüi ƒë·ªông th√†nh c√¥ng",
                "timestamp": "V·ª´a xong"
            }
        ]

def get_recent_foods():
    """L·∫•y m√≥n ƒÉn ƒë∆∞·ª£c t·∫°o g·∫ßn ƒë√¢y t·ª´ Firebase"""
    try:
        # L·∫•y d·ªØ li·ªáu foods t·ª´ Firebase
        foods_data = get_foods_data()

        # S·∫Øp x·∫øp theo th·ªùi gian t·∫°o v√† l·∫•y 5 m√≥n g·∫ßn nh·∫•t
        sorted_foods = sorted(foods_data, key=lambda x: x.get('created_at', ''), reverse=True)

        # Format l·∫°i d·ªØ li·ªáu ƒë·ªÉ hi·ªÉn th·ªã ƒë·∫πp h∆°n
        formatted_foods = []
        for food in sorted_foods[:5]:
            formatted_foods.append({
                "id": food.get('id', ''),
                "name": food.get('name', 'Kh√¥ng r√µ'),
                "nutrition": {
                    "calories": food.get('nutrition', {}).get('calories', 0)
                },
                "created_at": food.get('created_at', 'Kh√¥ng r√µ')
            })

        return formatted_foods
    except Exception as e:
        print(f"Error getting recent foods: {str(e)}")
        return []

def get_chart_data():
    """T·∫°o d·ªØ li·ªáu cho bi·ªÉu ƒë·ªì"""
    # D·ªØ li·ªáu ho·∫°t ƒë·ªông 7 ng√†y qua (gi·∫£ l·∫≠p)
    activity_labels = []
    activity_data = []
    
    for i in range(7):
        date = datetime.now() - timedelta(days=6-i)
        activity_labels.append(date.strftime("%d/%m"))
        activity_data.append(50 + i * 10 + (i % 3) * 20)  # Gi·∫£ l·∫≠p d·ªØ li·ªáu
    
    # D·ªØ li·ªáu lo·∫°i m√≥n ƒÉn
    food_type_labels = ["B·ªØa s√°ng", "B·ªØa tr∆∞a", "B·ªØa t·ªëi", "ƒê·ªì u·ªëng", "Tr√°ng mi·ªáng"]
    food_type_data = [25, 35, 30, 5, 5]  # Gi·∫£ l·∫≠p ph·∫ßn trƒÉm
    
    return {
        "activity_labels": activity_labels,
        "activity_data": activity_data,
        "food_type_labels": food_type_labels,
        "food_type_data": food_type_data
    }

def get_system_status():
    """Ki·ªÉm tra tr·∫°ng th√°i h·ªá th·ªëng"""
    try:
        # Ki·ªÉm tra AI service
        from groq_integration import groq_service
        ai_available = groq_service.available
        ai_type = "LLaMA 3 (Groq)" if ai_available else None
    except:
        ai_available = False
        ai_type = None
    
    # Ki·ªÉm tra Firebase
    try:
        firebase_connected = firestore_service.check_connection()
    except:
        firebase_connected = False
    
    return {
        "ai_available": ai_available,
        "ai_type": ai_type,
        "firebase_connected": firebase_connected
    }

@router.get("/", response_class=HTMLResponse)
@router.get("/dashboard", response_class=HTMLResponse)
async def admin_dashboard(
    request: Request,
    templates: Jinja2Templates = Depends(get_templates)
):
    """Trang dashboard admin"""
    # Ki·ªÉm tra x√°c th·ª±c admin
    admin_username = get_current_admin(request)
    if not admin_username:
        return RedirectResponse(url="/admin/login", status_code=302)
    try:
        # L·∫•y d·ªØ li·ªáu th·ªëng k√™
        stats = get_system_stats()
        recent_activities = get_recent_activities()
        recent_foods = get_recent_foods()
        chart_data = get_chart_data()
        system_status = get_system_status()
        
        return templates.TemplateResponse("admin/dashboard.html", {
            "request": request,
            "stats": stats,
            "recent_activities": recent_activities,
            "recent_foods": recent_foods,
            "activity_chart_labels": chart_data["activity_labels"],
            "activity_chart_data": chart_data["activity_data"],
            "food_type_labels": chart_data["food_type_labels"],
            "food_type_data": chart_data["food_type_data"],
            "system_status": system_status
        })
    except Exception as e:
        print(f"Error in admin dashboard: {str(e)}")
        # Tr·∫£ v·ªÅ trang v·ªõi d·ªØ li·ªáu m·∫∑c ƒë·ªãnh
        return templates.TemplateResponse("admin/dashboard.html", {
            "request": request,
            "stats": {"total_foods": 0, "active_users": 0, "total_meal_plans": 0, "api_calls_today": 0},
            "recent_activities": [],
            "recent_foods": [],
            "activity_chart_labels": [],
            "activity_chart_data": [],
            "food_type_labels": [],
            "food_type_data": [],
            "system_status": {"ai_available": False, "ai_type": None, "firebase_connected": False}
        })

@router.get("/fast-dashboard", response_class=HTMLResponse)
async def admin_fast_dashboard(
    request: Request,
    templates: Jinja2Templates = Depends(get_templates)
):
    """üöÄ Trang dashboard admin t·ªëi ∆∞u h√≥a"""
    # Ki·ªÉm tra x√°c th·ª±c admin
    admin_username = get_current_admin(request)
    if not admin_username:
        return RedirectResponse(url="/admin/fast-login", status_code=302)

    try:
        # L·∫•y d·ªØ li·ªáu th·ªëng k√™ (optimized - √≠t data h∆°n)
        stats = get_system_stats()
        recent_activities = get_recent_activities()[:5]  # Ch·ªâ l·∫•y 5 activities g·∫ßn nh·∫•t
        recent_foods = get_recent_foods()[:5]  # Ch·ªâ l·∫•y 5 foods g·∫ßn nh·∫•t
        system_status = get_system_status()

        return templates.TemplateResponse("admin/fast_dashboard.html", {
            "request": request,
            "stats": stats,
            "recent_activities": recent_activities,
            "recent_foods": recent_foods,
            "system_status": system_status
        })
    except Exception as e:
        print(f"Error in fast admin dashboard: {str(e)}")
        # Tr·∫£ v·ªÅ trang v·ªõi d·ªØ li·ªáu m·∫∑c ƒë·ªãnh
        return templates.TemplateResponse("admin/fast_dashboard.html", {
            "request": request,
            "stats": {"total_foods": 0, "active_users": 0, "total_meal_plans": 0, "api_calls_today": 0},
            "recent_activities": [],
            "recent_foods": [],
            "system_status": {"ai_available": False, "ai_type": None, "firebase_connected": False}
        })

@router.get("/users", response_class=HTMLResponse)
async def admin_users(
    request: Request,
    page: int = Query(1, ge=1),
    limit: int = Query(20, ge=1, le=100),
    search: Optional[str] = None,
    templates: Jinja2Templates = Depends(get_templates)
):
    """Trang qu·∫£n l√Ω ng∆∞·ªùi d√πng"""
    # Ki·ªÉm tra x√°c th·ª±c admin
    admin_username = get_current_admin(request)
    if not admin_username:
        return RedirectResponse(url="/admin/login", status_code=302)
    try:
        print(f"[ADMIN] Getting all users from Firebase...")
        # L·∫•y danh s√°ch ng∆∞·ªùi d√πng t·ª´ Firestore
        users = firestore_service.get_all_users()
        print(f"[ADMIN] Retrieved {len(users)} users from Firebase")

        # Debug: In ra m·ªôt v√†i user ƒë·∫ßu ti√™n
        if users:
            print(f"[ADMIN] First user sample: {users[0] if users else 'None'}")
            print(f"[ADMIN] User type: {type(users[0]) if users else 'None'}")
        else:
            print(f"[ADMIN] No users found!")
        
        # L·ªçc theo t·ª´ kh√≥a t√¨m ki·∫øm
        if search:
            search = search.lower()
            users = [
                user for user in users 
                if search in user.get('email', '').lower() or 
                   search in user.get('display_name', '').lower()
            ]
        
        # Ph√¢n trang
        total_users = len(users)
        start_idx = (page - 1) * limit
        end_idx = start_idx + limit
        users_page = users[start_idx:end_idx]
        
        # T√≠nh to√°n th√¥ng tin ph√¢n trang
        total_pages = (total_users + limit - 1) // limit
        has_prev = page > 1
        has_next = page < total_pages
        
        return templates.TemplateResponse("admin/users.html", {
            "request": request,
            "users": users_page,
            "current_page": page,
            "total_pages": total_pages,
            "total_users": total_users,
            "has_prev": has_prev,
            "has_next": has_next,
            "search": search or ""
        })
    except Exception as e:
        print(f"Error in admin users: {str(e)}")
        return templates.TemplateResponse("admin/users.html", {
            "request": request,
            "users": [],
            "current_page": 1,
            "total_pages": 1,
            "total_users": 0,
            "has_prev": False,
            "has_next": False,
            "search": search or "",
            "error": f"L·ªói khi t·∫£i d·ªØ li·ªáu ng∆∞·ªùi d√πng: {str(e)}"
        })

@router.get("/meal-plans", response_class=HTMLResponse)
async def admin_meal_plans(
    request: Request,
    page: int = Query(1, ge=1),
    limit: int = Query(20, ge=1, le=100),
    user_id: Optional[str] = None,
    templates: Jinja2Templates = Depends(get_templates)
):
    """Trang qu·∫£n l√Ω k·∫ø ho·∫°ch b·ªØa ƒÉn"""
    try:
        # L·∫•y danh s√°ch meal plans
        if user_id:
            meal_plans = firestore_service.get_user_meal_plans(user_id)
        else:
            meal_plans = firestore_service.get_all_meal_plans()
        
        # Ph√¢n trang
        total_plans = len(meal_plans)
        start_idx = (page - 1) * limit
        end_idx = start_idx + limit
        plans_page = meal_plans[start_idx:end_idx]
        
        # T√≠nh to√°n th√¥ng tin ph√¢n trang
        total_pages = (total_plans + limit - 1) // limit
        has_prev = page > 1
        has_next = page < total_pages
        
        return templates.TemplateResponse("admin/meal_plans.html", {
            "request": request,
            "meal_plans": plans_page,
            "current_page": page,
            "total_pages": total_pages,
            "total_plans": total_plans,
            "has_prev": has_prev,
            "has_next": has_next,
            "user_id": user_id or ""
        })
    except Exception as e:
        print(f"Error in admin meal plans: {str(e)}")
        return templates.TemplateResponse("admin/meal_plans.html", {
            "request": request,
            "meal_plans": [],
            "current_page": 1,
            "total_pages": 1,
            "total_plans": 0,
            "has_prev": False,
            "has_next": False,
            "user_id": user_id or "",
            "error": f"L·ªói khi t·∫£i d·ªØ li·ªáu k·∫ø ho·∫°ch b·ªØa ƒÉn: {str(e)}"
        })

@router.get("/foods", response_class=HTMLResponse)
async def admin_foods(
    request: Request,
    page: int = Query(1, ge=1),
    limit: int = Query(20, ge=1, le=100),
    search: Optional[str] = None,
    templates: Jinja2Templates = Depends(get_templates)
):
    """Trang qu·∫£n l√Ω m√≥n ƒÉn"""
    try:
        # L·∫•y danh s√°ch foods t·ª´ Firebase
        foods_data = get_foods_data()

        # L·ªçc theo t·ª´ kh√≥a t√¨m ki·∫øm
        if search:
            search = search.lower()
            foods_data = [
                food for food in foods_data
                if search in food.get('name', '').lower() or
                   search in food.get('description', '').lower()
            ]

        # Ph√¢n trang
        total_foods = len(foods_data)
        start_idx = (page - 1) * limit
        end_idx = start_idx + limit
        foods_page = foods_data[start_idx:end_idx]

        # T√≠nh to√°n th√¥ng tin ph√¢n trang
        total_pages = (total_foods + limit - 1) // limit
        has_prev = page > 1
        has_next = page < total_pages

        return templates.TemplateResponse("admin/foods.html", {
            "request": request,
            "foods": foods_page,
            "current_page": page,
            "total_pages": total_pages,
            "total_foods": total_foods,
            "has_prev": has_prev,
            "has_next": has_next,
            "search": search or ""
        })
    except Exception as e:
        print(f"Error in admin foods: {str(e)}")
        return templates.TemplateResponse("admin/foods.html", {
            "request": request,
            "foods": [],
            "current_page": 1,
            "total_pages": 1,
            "total_foods": 0,
            "has_prev": False,
            "has_next": False,
            "search": search or "",
            "error": f"L·ªói khi t·∫£i d·ªØ li·ªáu m√≥n ƒÉn: {str(e)}"
        })

@router.get("/reports", response_class=HTMLResponse)
async def admin_reports(
    request: Request,
    start_date: Optional[str] = Query(None),
    end_date: Optional[str] = Query(None),
    templates: Jinja2Templates = Depends(get_templates)
):
    """Trang b√°o c√°o v√† th·ªëng k√™"""
    try:
        # Thi·∫øt l·∫≠p ng√†y m·∫∑c ƒë·ªãnh n·∫øu kh√¥ng c√≥
        if not start_date:
            start_date = (datetime.now() - timedelta(days=30)).strftime("%Y-%m-%d")
        if not end_date:
            end_date = datetime.now().strftime("%Y-%m-%d")

        # L·∫•y d·ªØ li·ªáu metrics
        metrics = get_report_metrics(start_date, end_date)

        # L·∫•y d·ªØ li·ªáu bi·ªÉu ƒë·ªì
        chart_data = get_report_chart_data(start_date, end_date)

        # L·∫•y top users
        top_users = get_top_active_users()

        # L·∫•y l·ªói g·∫ßn ƒë√¢y
        recent_errors = get_recent_errors()

        return templates.TemplateResponse("admin/reports.html", {
            "request": request,
            "start_date": start_date,
            "end_date": end_date,
            "metrics": metrics,
            "activity_labels": chart_data["activity_labels"],
            "activity_data": chart_data["activity_data"],
            "api_calls_data": chart_data["api_calls_data"],
            "device_labels": chart_data["device_labels"],
            "device_data": chart_data["device_data"],
            "popular_foods_labels": chart_data["popular_foods_labels"],
            "popular_foods_data": chart_data["popular_foods_data"],
            "feature_labels": chart_data["feature_labels"],
            "feature_data": chart_data["feature_data"],
            "top_users": top_users,
            "recent_errors": recent_errors
        })
    except Exception as e:
        print(f"Error in admin reports: {str(e)}")
        # Tr·∫£ v·ªÅ trang v·ªõi d·ªØ li·ªáu m·∫∑c ƒë·ªãnh
        return templates.TemplateResponse("admin/reports.html", {
            "request": request,
            "start_date": start_date or datetime.now().strftime("%Y-%m-%d"),
            "end_date": end_date or datetime.now().strftime("%Y-%m-%d"),
            "metrics": {},
            "activity_labels": [],
            "activity_data": [],
            "api_calls_data": [],
            "device_labels": [],
            "device_data": [],
            "popular_foods_labels": [],
            "popular_foods_data": [],
            "feature_labels": [],
            "feature_data": [],
            "top_users": [],
            "recent_errors": [],
            "error": f"L·ªói khi t·∫£i b√°o c√°o: {str(e)}"
        })

def get_report_metrics(start_date: str, end_date: str):
    """L·∫•y c√°c metrics cho b√°o c√°o t·ª´ Firebase"""
    try:
        # L·∫•y d·ªØ li·ªáu th·∫≠t t·ª´ Firebase
        users = firestore_service.get_all_users()
        meal_plans = firestore_service.get_all_meal_plans()

        # T√≠nh to√°n metrics th·∫≠t
        total_users = len(users)
        total_meal_plans = len(meal_plans)

        # ƒê·∫øm users m·ªõi trong kho·∫£ng th·ªùi gian
        new_users = 0
        for user in users:
            created_at = user.get('created_at')
            if created_at:
                # Chuy·ªÉn ƒë·ªïi created_at th√†nh string n·∫øu c·∫ßn
                if hasattr(created_at, 'strftime'):
                    created_at_str = created_at.strftime('%Y-%m-%d')
                else:
                    created_at_str = str(created_at)[:10]  # L·∫•y ph·∫ßn YYYY-MM-DD

                if start_date <= created_at_str <= end_date:
                    new_users += 1

        # ƒê·∫øm meal plans m·ªõi trong kho·∫£ng th·ªùi gian
        new_meal_plans = 0
        for plan in meal_plans:
            created_at = plan.get('created_at')
            if created_at:
                # Chuy·ªÉn ƒë·ªïi created_at th√†nh string n·∫øu c·∫ßn
                if hasattr(created_at, 'strftime'):
                    created_at_str = created_at.strftime('%Y-%m-%d')
                else:
                    created_at_str = str(created_at)[:10]  # L·∫•y ph·∫ßn YYYY-MM-DD

                if start_date <= created_at_str <= end_date:
                    new_meal_plans += 1

        # T√≠nh to√°n t·ª∑ l·ªá tƒÉng tr∆∞·ªüng (gi·∫£ ƒë·ªãnh)
        users_growth = (new_users / max(total_users - new_users, 1)) * 100 if total_users > 0 else 0
        meal_plans_growth = (new_meal_plans / max(total_meal_plans - new_meal_plans, 1)) * 100 if total_meal_plans > 0 else 0

        return {
            "total_api_calls": total_users * 50 + total_meal_plans * 20,  # ∆Ø·ªõc t√≠nh
            "api_calls_growth": min(users_growth + meal_plans_growth, 50),  # Gi·ªõi h·∫°n 50%
            "new_users": new_users,
            "new_users_growth": users_growth,
            "meal_plans_created": new_meal_plans,
            "meal_plans_growth": meal_plans_growth,
            "activity_rate": min((total_users * 0.7), 100),  # Gi·∫£ ƒë·ªãnh 70% users active
            "activity_rate_change": min(users_growth * 0.5, 10)  # Thay ƒë·ªïi activity d·ª±a tr√™n user growth
        }
    except Exception as e:
        print(f"Error getting report metrics: {str(e)}")
        return {
            "total_api_calls": 0,
            "api_calls_growth": 0,
            "new_users": 0,
            "new_users_growth": 0,
            "meal_plans_created": 0,
            "meal_plans_growth": 0,
            "activity_rate": 0,
            "activity_rate_change": 0
        }

def get_report_chart_data(start_date: str, end_date: str):
    """L·∫•y d·ªØ li·ªáu cho c√°c bi·ªÉu ƒë·ªì t·ª´ Firebase"""
    try:
        # L·∫•y d·ªØ li·ªáu th·∫≠t t·ª´ Firebase
        users = firestore_service.get_all_users()
        meal_plans = firestore_service.get_all_meal_plans()

        # T·∫°o d·ªØ li·ªáu cho bi·ªÉu ƒë·ªì ho·∫°t ƒë·ªông theo ng√†y
        days = []
        activity_data = []
        api_calls_data = []

        # T·∫°o d·ªØ li·ªáu cho 30 ng√†y g·∫ßn ƒë√¢y
        for i in range(30):
            date = datetime.now() - timedelta(days=29-i)
            days.append(date.strftime("%d/%m"))

            # ƒê·∫øm ho·∫°t ƒë·ªông trong ng√†y (gi·∫£ l·∫≠p d·ª±a tr√™n d·ªØ li·ªáu th·∫≠t)
            target_date = date.strftime("%Y-%m-%d")

            # ƒê·∫øm users ƒë∆∞·ª£c t·∫°o trong ng√†y
            day_users = 0
            for u in users:
                created_at = u.get('created_at')
                if created_at:
                    if hasattr(created_at, 'strftime'):
                        created_date = created_at.strftime('%Y-%m-%d')
                    else:
                        created_date = str(created_at)[:10]
                    if created_date == target_date:
                        day_users += 1

            # ƒê·∫øm meal plans ƒë∆∞·ª£c t·∫°o trong ng√†y
            day_plans = 0
            for p in meal_plans:
                created_at = p.get('created_at')
                if created_at:
                    if hasattr(created_at, 'strftime'):
                        created_date = created_at.strftime('%Y-%m-%d')
                    else:
                        created_date = str(created_at)[:10]
                    if created_date == target_date:
                        day_plans += 1

            day_activity = day_users * 5
            day_api_calls = day_plans * 10

            activity_data.append(max(day_activity, 10 + i % 20))  # ƒê·∫£m b·∫£o c√≥ d·ªØ li·ªáu hi·ªÉn th·ªã
            api_calls_data.append(max(day_api_calls, 50 + i % 30))

        # Th·ªëng k√™ m√≥n ƒÉn ph·ªï bi·∫øn t·ª´ Firebase
        foods_data = get_foods_data()
        popular_foods = {}
        for food in foods_data[:20]:  # L·∫•y 20 m√≥n ƒë·∫ßu ti√™n
            name = food.get('name', 'Unknown')
            # Gi·∫£ l·∫≠p popularity d·ª±a tr√™n calories (m√≥n c√≥ calories cao th∆∞·ªùng ƒë∆∞·ª£c quan t√¢m)
            calories = food.get('nutrition', {}).get('calories', 0)
            popularity = min(calories / 10, 50)  # Chuy·ªÉn ƒë·ªïi calories th√†nh popularity score
            popular_foods[name] = popularity

        # S·∫Øp x·∫øp v√† l·∫•y top 5
        sorted_foods = sorted(popular_foods.items(), key=lambda x: x[1], reverse=True)[:5]
        popular_foods_labels = [item[0] for item in sorted_foods]
        popular_foods_data = [int(item[1]) for item in sorted_foods]

        return {
            "activity_labels": days,
            "activity_data": activity_data,
            "api_calls_data": api_calls_data,
            "device_labels": ["Mobile", "Desktop", "Tablet"],
            "device_data": [65, 25, 10],  # Gi·∫£ l·∫≠p device usage
            "popular_foods_labels": popular_foods_labels,
            "popular_foods_data": popular_foods_data,
            "feature_labels": ["T·∫°o meal plan", "T√¨m ki·∫øm th·ª±c ph·∫©m", "Chat AI", "Theo d√µi dinh d∆∞·ª°ng", "B√°o c√°o"],
            "feature_data": [len(meal_plans), len(foods_data)//10, len(users)//2, len(users)//3, len(users)//5]
        }
    except Exception as e:
        print(f"Error getting chart data: {str(e)}")
        return {
            "activity_labels": [],
            "activity_data": [],
            "api_calls_data": [],
            "device_labels": [],
            "device_data": [],
            "popular_foods_labels": [],
            "popular_foods_data": [],
            "feature_labels": [],
            "feature_data": []
        }

def get_top_active_users():
    """L·∫•y danh s√°ch ng∆∞·ªùi d√πng ho·∫°t ƒë·ªông nh·∫•t t·ª´ Firebase"""
    try:
        # L·∫•y d·ªØ li·ªáu th·∫≠t t·ª´ Firebase
        users = firestore_service.get_all_users()
        meal_plans = firestore_service.get_all_meal_plans()

        # T√≠nh to√°n activity cho m·ªói user
        user_activities = {}

        # ƒê·∫øm meal plans cho m·ªói user
        for plan in meal_plans:
            user_id = plan.get('user_id', '')
            if user_id:
                if user_id not in user_activities:
                    user_activities[user_id] = {'meal_plans_count': 0, 'activity_count': 0}
                user_activities[user_id]['meal_plans_count'] += 1
                user_activities[user_id]['activity_count'] += 5  # M·ªói meal plan = 5 ƒëi·ªÉm activity

        # T·∫°o danh s√°ch top users
        top_users = []
        for user in users[:10]:  # L·∫•y t·ªëi ƒëa 10 users
            user_id = user.get('uid', user.get('id', ''))
            activity_data = user_activities.get(user_id, {'meal_plans_count': 0, 'activity_count': 0})

            # T√≠nh th√™m activity t·ª´ th√¥ng tin user
            if user.get('last_login'):
                activity_data['activity_count'] += 10
            if user.get('profile_completed'):
                activity_data['activity_count'] += 5

            top_users.append({
                "display_name": user.get('display_name', user.get('name', 'Ng∆∞·ªùi d√πng ·∫©n danh')),
                "email": user.get('email', 'Kh√¥ng c√≥ email'),
                "photo_url": user.get('photo_url'),
                "activity_count": activity_data['activity_count'],
                "meal_plans_count": activity_data['meal_plans_count'],
                "last_activity": user.get('last_login', user.get('created_at', 'Kh√¥ng r√µ'))
            })

        # S·∫Øp x·∫øp theo activity_count v√† l·∫•y top 5
        top_users.sort(key=lambda x: x['activity_count'], reverse=True)
        return top_users[:5]

    except Exception as e:
        print(f"Error getting top users: {str(e)}")
        return [
            {
                "display_name": "Kh√¥ng c√≥ d·ªØ li·ªáu",
                "email": "system@openfood.com",
                "photo_url": None,
                "activity_count": 0,
                "meal_plans_count": 0,
                "last_activity": "Kh√¥ng r√µ"
            }
        ]

def get_recent_errors():
    """L·∫•y danh s√°ch l·ªói g·∫ßn ƒë√¢y"""
    try:
        # Gi·∫£ l·∫≠p d·ªØ li·ªáu l·ªói
        return [
            {
                "type": "API Timeout",
                "message": "Groq API timeout khi t·∫°o meal plan",
                "level": "warning",
                "count": 3
            },
            {
                "type": "Database Connection",
                "message": "K·∫øt n·ªëi Firestore b·ªã gi√°n ƒëo·∫°n",
                "level": "error",
                "count": 1
            }
        ]
    except Exception as e:
        print(f"Error getting recent errors: {str(e)}")
        return []

@router.get("/settings", response_class=HTMLResponse)
async def admin_settings(
    request: Request,
    templates: Jinja2Templates = Depends(get_templates)
):
    """Trang c·∫•u h√¨nh h·ªá th·ªëng"""
    try:
        # L·∫•y tr·∫°ng th√°i h·ªá th·ªëng
        system_status = get_system_status()

        # L·∫•y c·∫•u h√¨nh hi·ªán t·∫°i
        settings = get_current_settings()

        return templates.TemplateResponse("admin/settings.html", {
            "request": request,
            "system_status": system_status,
            "settings": settings
        })
    except Exception as e:
        print(f"Error in admin settings: {str(e)}")
        return templates.TemplateResponse("admin/settings.html", {
            "request": request,
            "system_status": {
                "database_connected": False,
                "ai_service_available": False,
                "firebase_connected": False,
                "storage_available": False
            },
            "settings": {},
            "error": f"L·ªói khi t·∫£i c·∫•u h√¨nh: {str(e)}"
        })

def get_current_settings():
    """L·∫•y c·∫•u h√¨nh hi·ªán t·∫°i c·ªßa h·ªá th·ªëng"""
    try:
        import os

        # L·∫•y c·∫•u h√¨nh t·ª´ bi·∫øn m√¥i tr∆∞·ªùng v√† c·∫•u h√¨nh m·∫∑c ƒë·ªãnh
        settings = {
            # AI & API Settings
            "groq_api_key": "***" if os.getenv("GROQ_API_KEY") else "",
            "groq_model": "llama3-8b-8192",
            "groq_timeout": 30,
            "usda_api_key": "***" if os.getenv("USDA_API_KEY") else "",
            "usda_max_results": 10,
            "usda_cache_enabled": True,
            "rate_limit_per_minute": 60,
            "rate_limit_per_day": 1000,
            "cache_expiry_hours": 24,
            "cache_enabled": True,

            # Database Settings
            "firebase_project_id": os.getenv("FIREBASE_PROJECT_ID", ""),
            "firebase_storage_bucket": os.getenv("FIREBASE_STORAGE_BUCKET", ""),
            "firebase_emulator_enabled": False,
            "connection_pool_size": 10,
            "query_timeout": 30,
            "enable_query_logging": False,

            # Security Settings
            "jwt_secret": "***" if os.getenv("JWT_SECRET") else "",
            "token_expiry_hours": 24,
            "require_email_verification": False,
            "allowed_origins": "*",
            "enable_cors": True,
            "force_https": False,

            # Performance Settings
            "max_workers": 4,
            "request_timeout": 30,
            "enable_gzip": True,
            "log_level": "INFO",
            "enable_metrics": True,
            "enable_health_check": True,

            # Notification Settings
            "smtp_host": "",
            "smtp_port": 587,
            "smtp_username": "",
            "smtp_password": "",
            "admin_email": "",
            "alert_on_errors": True,
            "alert_on_high_usage": True,
            "daily_reports": False
        }

        return settings
    except Exception as e:
        print(f"Error getting current settings: {str(e)}")
        return {}

# API endpoints for food records CRUD
@router.get("/api/foods/{food_id}")
async def get_food_api(food_id: str):
    """API ƒë·ªÉ l·∫•y th√¥ng tin m·ªôt food record"""
    try:
        print(f"[API] Getting food record with ID: {food_id}")
        food = firestore_service.get_food_record(food_id)
        print(f"[API] Food record result: {food is not None}")
        if food:
            print(f"[API] Food record data keys: {list(food.keys()) if food else 'None'}")
            return {"success": True, "food": food}
        else:
            print(f"[API] Food record not found for ID: {food_id}")
            return {"success": False, "message": "Kh√¥ng t√¨m th·∫•y food record"}
    except Exception as e:
        print(f"[API] Error getting food record: {str(e)}")
        return {"success": False, "message": f"L·ªói: {str(e)}"}

# API endpoints for meal plans CRUD
@router.get("/api/meal-plans/{plan_id}")
async def get_meal_plan_api(plan_id: str):
    """API ƒë·ªÉ l·∫•y th√¥ng tin m·ªôt meal plan"""
    try:
        print(f"[API] Getting meal plan with ID: {plan_id}")
        meal_plan_dict = firestore_service.get_meal_plan_dict(plan_id)
        print(f"[API] Meal plan result: {meal_plan_dict is not None}")
        if meal_plan_dict:
            print(f"[API] Meal plan data keys: {list(meal_plan_dict.keys())}")
            return {"success": True, "meal_plan": meal_plan_dict}
        else:
            print(f"[API] Meal plan not found for ID: {plan_id}")
            return {"success": False, "message": "Kh√¥ng t√¨m th·∫•y meal plan"}
    except Exception as e:
        print(f"[API] Error getting meal plan: {str(e)}")
        return {"success": False, "message": f"L·ªói: {str(e)}"}

@router.put("/api/meal-plans/{plan_id}")
async def update_meal_plan_api(plan_id: str, meal_plan_data: dict):
    """API ƒë·ªÉ c·∫≠p nh·∫≠t meal plan"""
    try:
        print(f"[API] Updating meal plan with ID: {plan_id}")
        print(f"[API] Update data: {meal_plan_data}")

        # C·∫≠p nh·∫≠t meal plan trong Firebase
        success = firestore_service.update_meal_plan(plan_id, meal_plan_data)

        if success:
            print(f"[API] Meal plan updated successfully")
            return {"success": True, "message": "C·∫≠p nh·∫≠t meal plan th√†nh c√¥ng"}
        else:
            print(f"[API] Failed to update meal plan")
            return {"success": False, "message": "Kh√¥ng th·ªÉ c·∫≠p nh·∫≠t meal plan"}
    except Exception as e:
        print(f"[API] Error updating meal plan: {str(e)}")
        return {"success": False, "message": f"L·ªói: {str(e)}"}

@router.delete("/api/meal-plans/{plan_id}")
async def delete_meal_plan_api(plan_id: str):
    """API ƒë·ªÉ x√≥a meal plan"""
    try:
        print(f"[API] Deleting meal plan with ID: {plan_id}")

        # X√≥a meal plan t·ª´ Firebase
        success = firestore_service.delete_meal_plan(plan_id)

        if success:
            print(f"[API] Meal plan deleted successfully")
            return {"success": True, "message": "X√≥a meal plan th√†nh c√¥ng"}
        else:
            print(f"[API] Failed to delete meal plan")
            return {"success": False, "message": "Kh√¥ng th·ªÉ x√≥a meal plan"}
    except Exception as e:
        print(f"[API] Error deleting meal plan: {str(e)}")
        return {"success": False, "message": f"L·ªói: {str(e)}"}

# API endpoints for users
@router.get("/api/users/{user_id}")
async def get_user_api(user_id: str):
    """API ƒë·ªÉ l·∫•y th√¥ng tin m·ªôt user"""
    try:
        print(f"[API] Getting user with ID: {user_id}")
        user = firestore_service.get_user_by_id(user_id)
        print(f"[API] User result: {user is not None}")
        if user:
            print(f"[API] User data keys: {list(user.keys()) if isinstance(user, dict) else 'Not a dict'}")
            return {"success": True, "user": user}
        else:
            print(f"[API] User not found for ID: {user_id}")
            return {"success": False, "message": "Kh√¥ng t√¨m th·∫•y ng∆∞·ªùi d√πng"}
    except Exception as e:
        print(f"[API] Error getting user: {str(e)}")
        return {"success": False, "message": f"L·ªói: {str(e)}"}

@router.delete("/api/users/{user_id}")
async def delete_user_api(user_id: str, request: Request):
    """API ƒë·ªÉ x√≥a user v√† t·∫•t c·∫£ d·ªØ li·ªáu li√™n quan"""
    try:
        # Ki·ªÉm tra x√°c th·ª±c admin
        admin_username = get_current_admin(request)
        if not admin_username:
            return {"success": False, "message": "Kh√¥ng c√≥ quy·ªÅn truy c·∫≠p"}

        print(f"[API] Admin {admin_username} deleting user: {user_id}")

        # Ki·ªÉm tra user c√≥ t·ªìn t·∫°i kh√¥ng
        user = firestore_service.get_user_by_id(user_id)
        if not user:
            return {"success": False, "message": "Kh√¥ng t√¨m th·∫•y ng∆∞·ªùi d√πng"}

        # X√≥a user v√† t·∫•t c·∫£ d·ªØ li·ªáu li√™n quan
        success = firestore_service.delete_user(user_id)

        if success:
            print(f"[API] Successfully deleted user: {user_id}")
            return {
                "success": True,
                "message": f"ƒê√£ x√≥a ng∆∞·ªùi d√πng {user.get('name', user.get('email', user_id))} v√† t·∫•t c·∫£ d·ªØ li·ªáu li√™n quan"
            }
        else:
            print(f"[API] Failed to delete user: {user_id}")
            return {"success": False, "message": "Kh√¥ng th·ªÉ x√≥a ng∆∞·ªùi d√πng"}

    except Exception as e:
        print(f"[API] Error deleting user: {str(e)}")
        return {"success": False, "message": f"L·ªói: {str(e)}"}

@router.get("/api/foods/{food_id}/test")
async def test_food_api(food_id: str):
    """Test API ƒë·ªÉ debug"""
    return {
        "success": True,
        "food_id": food_id,
        "message": "Test API ho·∫°t ƒë·ªông b√¨nh th∆∞·ªùng",
        "food": {
            "id": food_id,
            "name": "Test Food",
            "description": "Test Description",
            "mealType": "Test Meal",
            "nutrition": {
                "calories": 100,
                "protein": 10,
                "fat": 5,
                "carbs": 15
            }
        }
    }

@router.post("/api/foods")
async def create_food_api(food_data: dict):
    """API ƒë·ªÉ t·∫°o food record m·ªõi (ch·ªâ cho demo - th·ª±c t·∫ø food_records ƒë∆∞·ª£c t·∫°o t·ª´ app)"""
    try:
        return {"success": False, "message": "T√≠nh nƒÉng t·∫°o food record m·ªõi ch∆∞a ƒë∆∞·ª£c h·ªó tr·ª£. Food records ƒë∆∞·ª£c t·∫°o t·ª´ ·ª©ng d·ª•ng mobile."}
    except Exception as e:
        return {"success": False, "message": f"L·ªói: {str(e)}"}

@router.put("/api/foods/{food_id}")
async def update_food_api(food_id: str, food_data: dict):
    """API ƒë·ªÉ c·∫≠p nh·∫≠t food record"""
    try:
        success = firestore_service.update_food_record(food_id, food_data)
        if success:
            return {"success": True, "message": "C·∫≠p nh·∫≠t food record th√†nh c√¥ng"}
        else:
            return {"success": False, "message": "Kh√¥ng th·ªÉ c·∫≠p nh·∫≠t food record"}
    except Exception as e:
        return {"success": False, "message": f"L·ªói: {str(e)}"}

@router.delete("/api/foods/{food_id}")
async def delete_food_api(food_id: str):
    """API ƒë·ªÉ x√≥a food record"""
    try:
        success = firestore_service.delete_food_record(food_id)
        if success:
            return {"success": True, "message": "X√≥a food record th√†nh c√¥ng"}
        else:
            return {"success": False, "message": "Kh√¥ng th·ªÉ x√≥a food record"}
    except Exception as e:
        return {"success": False, "message": f"L·ªói: {str(e)}"}

# üöÄ FAST API ENDPOINTS FOR OPTIMIZED ADMIN

@router.get("/api/quick-stats")
async def get_quick_stats(request: Request):
    """üöÄ API l·∫•y stats nhanh cho fast dashboard"""
    # Ki·ªÉm tra x√°c th·ª±c admin
    admin_username = get_current_admin(request)
    if not admin_username:
        return {"success": False, "message": "Unauthorized"}

    try:
        stats = get_system_stats()
        return {"success": True, "data": stats}
    except Exception as e:
        return {"success": False, "message": f"Error: {str(e)}"}

@router.get("/api/quick-export")
async def quick_export(request: Request):
    """üöÄ API xu·∫•t b√°o c√°o nhanh (CSV ƒë∆°n gi·∫£n)"""
    # Ki·ªÉm tra x√°c th·ª±c admin
    admin_username = get_current_admin(request)
    if not admin_username:
        return {"success": False, "message": "Unauthorized"}

    try:
        # T·∫°o CSV ƒë∆°n gi·∫£n v·ªõi d·ªØ li·ªáu c∆° b·∫£n
        stats = get_system_stats()
        csv_content = f"""Metric,Value
Total Foods,{stats.get('total_foods', 0)}
Active Users,{stats.get('active_users', 0)}
Total Meal Plans,{stats.get('total_meal_plans', 0)}
API Calls Today,{stats.get('api_calls_today', 0)}
Export Time,{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
"""

        return Response(
            content=csv_content,
            media_type="text/csv",
            headers={"Content-Disposition": f"attachment; filename=quick_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv"}
        )
    except Exception as e:
        return {"success": False, "message": f"Error: {str(e)}"}

@router.post("/api/fast-refresh")
async def fast_refresh_data(request: Request):
    """üöÄ API l√†m m·ªõi d·ªØ li·ªáu nhanh"""
    # Ki·ªÉm tra x√°c th·ª±c admin
    admin_username = get_current_admin(request)
    if not admin_username:
        return {"success": False, "message": "Unauthorized"}

    try:
        # Clear cache n·∫øu c√≥
        # L·∫•y d·ªØ li·ªáu m·ªõi
        stats = get_system_stats()
        recent_activities = get_recent_activities()[:3]  # Ch·ªâ l·∫•y 3 activities

        return {
            "success": True,
            "data": {
                "stats": stats,
                "recent_activities": recent_activities,
                "timestamp": datetime.now().isoformat()
            }
        }
    except Exception as e:
        return {"success": False, "message": f"Error: {str(e)}"}

# API endpoint t·∫°o download token
@router.post("/api/create-download-token")
async def create_download_token(request: Request):
    """T·∫°o token t·∫°m th·ªùi ƒë·ªÉ download b√°o c√°o"""
    # Ki·ªÉm tra x√°c th·ª±c admin
    admin_username = get_current_admin(request)
    if not admin_username:
        raise HTTPException(status_code=401, detail="Unauthorized")

    # T·∫°o token ng·∫´u nhi√™n
    token = secrets.token_urlsafe(32)

    # L∆∞u token v·ªõi th·ªùi gian h·∫øt h·∫°n (5 ph√∫t)
    download_tokens[token] = {
        "admin": admin_username,
        "created_at": time.time(),
        "expires_at": time.time() + 300  # 5 ph√∫t
    }

    # X√≥a c√°c token ƒë√£ h·∫øt h·∫°n
    current_time = time.time()
    expired_tokens = [t for t, data in download_tokens.items() if data["expires_at"] < current_time]
    for expired_token in expired_tokens:
        del download_tokens[expired_token]

    return {"success": True, "token": token}

# API endpoint debug dependencies
@router.get("/api/debug/dependencies")
async def debug_dependencies(request: Request):
    """Debug endpoint ƒë·ªÉ ki·ªÉm tra dependencies"""
    admin_username = get_current_admin(request)
    if not admin_username:
        raise HTTPException(status_code=401, detail="Unauthorized")

    return {
        "excel_available": EXCEL_AVAILABLE,
        "word_available": WORD_AVAILABLE,
        "python_version": __import__('sys').version,
        "dependencies": {
            "openpyxl": "Available" if EXCEL_AVAILABLE else "Not installed",
            "python-docx": "Available" if WORD_AVAILABLE else "Not installed"
        }
    }

# API endpoint test download ƒë∆°n gi·∫£n
@router.get("/api/test-download")
async def test_download(request: Request):
    """Test download ƒë∆°n gi·∫£n"""
    admin_username = get_current_admin(request)
    if not admin_username:
        raise HTTPException(status_code=401, detail="Unauthorized")

    # T·∫°o file text ƒë∆°n gi·∫£n
    content = f"""ADMIN REPORT TEST
Admin: {admin_username}
Date: {datetime.now().isoformat()}
Status: Working!

This is a test download to verify the download functionality.
"""

    def generate_content():
        yield content

    return StreamingResponse(
        generate_content(),
        media_type="text/plain",
        headers={"Content-Disposition": f"attachment; filename=test_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt"}
    )

# API endpoint test CSV
@router.get("/api/test-csv")
async def test_csv(request: Request):
    """Test CSV export ƒë∆°n gi·∫£n"""
    admin_username = get_current_admin(request)
    if not admin_username:
        raise HTTPException(status_code=401, detail="Unauthorized")

    # T·∫°o CSV ƒë∆°n gi·∫£n
    csv_content = f"""ADMIN REPORT TEST
Admin,{admin_username}
Date,{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
Status,Working

Sample Data
Name,Email,Age
John Doe,john@example.com,25
Jane Smith,jane@example.com,30
Bob Johnson,bob@example.com,35
"""

    def generate_csv():
        yield csv_content

    return StreamingResponse(
        generate_csv(),
        media_type="text/csv",
        headers={"Content-Disposition": f"attachment; filename=test_csv_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv"}
    )

# API endpoint xu·∫•t b√°o c√°o
@router.get("/api/export/report")
async def export_report(
    request: Request,
    format: str = Query("csv", description="Format: csv, json, excel, word"),
    start_date: Optional[str] = Query(None),
    end_date: Optional[str] = Query(None),
    token: Optional[str] = Query(None, description="Download token")
):
    """API ƒë·ªÉ xu·∫•t b√°o c√°o d·ªØ li·ªáu"""
    # Ki·ªÉm tra x√°c th·ª±c admin (session ho·∫∑c token)
    admin_username = get_current_admin(request)

    # N·∫øu kh√¥ng c√≥ session, ki·ªÉm tra token
    if not admin_username and token:
        if token in download_tokens:
            token_data = download_tokens[token]
            current_time = time.time()

            # Ki·ªÉm tra token c√≤n h·∫°n kh√¥ng
            if token_data["expires_at"] > current_time:
                admin_username = token_data["admin"]
                # X√≥a token sau khi s·ª≠ d·ª•ng (one-time use)
                del download_tokens[token]
            else:
                # Token ƒë√£ h·∫øt h·∫°n
                del download_tokens[token]
                raise HTTPException(status_code=401, detail="Token expired")
        else:
            raise HTTPException(status_code=401, detail="Invalid token")

    if not admin_username:
        raise HTTPException(status_code=401, detail="Unauthorized")

    try:
        # Thi·∫øt l·∫≠p ng√†y m·∫∑c ƒë·ªãnh
        if not start_date:
            start_date = (datetime.now() - timedelta(days=30)).strftime("%Y-%m-%d")
        if not end_date:
            end_date = datetime.now().strftime("%Y-%m-%d")

        print(f"[EXPORT] Starting export for admin: {admin_username}, format: {format}")

        # L·∫•y d·ªØ li·ªáu t·ª´ Firebase v·ªõi error handling
        try:
            users = firestore_service.get_all_users()
            print(f"[EXPORT] Got {len(users)} users")
        except Exception as e:
            print(f"[EXPORT] Error getting users: {e}")
            users = []

        try:
            food_records = firestore_service.get_all_food_records()
            print(f"[EXPORT] Got {len(food_records)} food records")
        except Exception as e:
            print(f"[EXPORT] Error getting food records: {e}")
            food_records = []

        try:
            meal_plans = firestore_service.get_all_meal_plans()
            print(f"[EXPORT] Got {len(meal_plans)} meal plans")
        except Exception as e:
            print(f"[EXPORT] Error getting meal plans: {e}")
            meal_plans = []

        # T·∫°o b√°o c√°o t·ªïng h·ª£p ƒë∆°n gi·∫£n
        report_data = {
            "export_date": datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
            "period": f"{start_date} to {end_date}",
            "admin": admin_username,
            "summary": {
                "total_users": len(users),
                "total_food_records": len(food_records),
                "total_meal_plans": len(meal_plans),
                "active_users": len([u for u in users if u.get('last_login_at')]) if users else 0,
            },
            "users": users[:10],  # Ch·ªâ 10 users ƒë·ªÉ test
            "recent_food_records": food_records[:10],  # 10 food records
            "recent_meal_plans": meal_plans[:5]  # 5 meal plans
        }

        if format.lower() == "json":
            # Xu·∫•t JSON
            json_str = json.dumps(report_data, indent=2, ensure_ascii=False, default=str)

            def generate_json():
                yield json_str

            return StreamingResponse(
                generate_json(),
                media_type="application/json",
                headers={"Content-Disposition": f"attachment; filename=admin_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"}
            )

        elif format.lower() == "excel" and EXCEL_AVAILABLE:
            # Xu·∫•t Excel
            wb = Workbook()

            # X√≥a sheet m·∫∑c ƒë·ªãnh
            wb.remove(wb.active)

            # T·∫°o sheet Summary
            summary_ws = wb.create_sheet("Summary")
            summary_ws.append(["ADMIN REPORT SUMMARY"])
            summary_ws.append([])
            summary_ws.append(["Export Date", report_data['export_date']])
            summary_ws.append(["Period", report_data['period']])
            summary_ws.append(["Admin", report_data['admin']])
            summary_ws.append([])

            # Th√™m summary data
            summary_ws.append(["STATISTICS"])
            for key, value in report_data['summary'].items():
                summary_ws.append([key.replace('_', ' ').title(), value])

            # Format header
            for row in summary_ws.iter_rows(min_row=1, max_row=1):
                for cell in row:
                    cell.font = Font(bold=True, size=14)
                    cell.fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
                    cell.font = Font(bold=True, color="FFFFFF")

            # T·∫°o sheet Users
            if users:
                users_ws = wb.create_sheet("Users")
                user_headers = ['ID', 'Email', 'Name', 'Created At', 'Last Login', 'Age', 'Gender']
                users_ws.append(user_headers)

                for user in users[:100]:
                    row_data = [
                        user.get('id', ''),
                        user.get('email', ''),
                        user.get('name', ''),
                        str(user.get('created_at', '')),
                        str(user.get('last_login_at', '')),
                        user.get('age', ''),
                        user.get('gender', '')
                    ]
                    users_ws.append(row_data)

                # Format headers
                for cell in users_ws[1]:
                    cell.font = Font(bold=True)
                    cell.fill = PatternFill(start_color="D9E1F2", end_color="D9E1F2", fill_type="solid")

            # T·∫°o sheet Food Records
            if food_records:
                foods_ws = wb.create_sheet("Food Records")
                food_headers = ['ID', 'User ID', 'Description', 'Meal Type', 'Date', 'Created At']
                foods_ws.append(food_headers)

                for food in food_records[:50]:
                    row_data = [
                        food.get('id', ''),
                        food.get('user_id', ''),
                        food.get('description', ''),
                        food.get('mealType', ''),
                        str(food.get('date', '')),
                        str(food.get('created_at', ''))
                    ]
                    foods_ws.append(row_data)

                # Format headers
                for cell in foods_ws[1]:
                    cell.font = Font(bold=True)
                    cell.fill = PatternFill(start_color="E2EFDA", end_color="E2EFDA", fill_type="solid")

            # L∆∞u Excel file
            excel_buffer = io.BytesIO()
            wb.save(excel_buffer)
            excel_buffer.seek(0)

            def generate_excel():
                yield excel_buffer.getvalue()

            return StreamingResponse(
                generate_excel(),
                media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                headers={"Content-Disposition": f"attachment; filename=admin_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"}
            )

        elif format.lower() == "word" and WORD_AVAILABLE:
            # Xu·∫•t Word
            doc = Document()

            # Title
            title = doc.add_heading('ADMIN REPORT', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Report info
            doc.add_heading('Report Information', level=1)
            info_table = doc.add_table(rows=3, cols=2)
            info_table.style = 'Table Grid'

            info_table.cell(0, 0).text = 'Export Date'
            info_table.cell(0, 1).text = report_data['export_date']
            info_table.cell(1, 0).text = 'Period'
            info_table.cell(1, 1).text = report_data['period']
            info_table.cell(2, 0).text = 'Admin'
            info_table.cell(2, 1).text = report_data['admin']

            # Summary
            doc.add_heading('Summary Statistics', level=1)
            summary_table = doc.add_table(rows=len(report_data['summary']) + 1, cols=2)
            summary_table.style = 'Table Grid'

            # Headers
            summary_table.cell(0, 0).text = 'Metric'
            summary_table.cell(0, 1).text = 'Value'

            # Data
            for i, (key, value) in enumerate(report_data['summary'].items(), 1):
                summary_table.cell(i, 0).text = key.replace('_', ' ').title()
                summary_table.cell(i, 1).text = str(value)

            # Users section
            if users:
                doc.add_heading('Recent Users (Top 20)', level=1)
                users_table = doc.add_table(rows=min(21, len(users) + 1), cols=4)
                users_table.style = 'Table Grid'

                # Headers
                users_table.cell(0, 0).text = 'Email'
                users_table.cell(0, 1).text = 'Name'
                users_table.cell(0, 2).text = 'Created At'
                users_table.cell(0, 3).text = 'Gender'

                # Data (top 20)
                for i, user in enumerate(users[:20], 1):
                    users_table.cell(i, 0).text = user.get('email', '')[:30]  # Limit length
                    users_table.cell(i, 1).text = user.get('name', '')[:20]
                    users_table.cell(i, 2).text = str(user.get('created_at', ''))[:10]
                    users_table.cell(i, 3).text = user.get('gender', '')

            # Food records section
            if food_records:
                doc.add_heading('Recent Food Records (Top 15)', level=1)
                foods_table = doc.add_table(rows=min(16, len(food_records) + 1), cols=4)
                foods_table.style = 'Table Grid'

                # Headers
                foods_table.cell(0, 0).text = 'User ID'
                foods_table.cell(0, 1).text = 'Description'
                foods_table.cell(0, 2).text = 'Meal Type'
                foods_table.cell(0, 3).text = 'Date'

                # Data (top 15)
                for i, food in enumerate(food_records[:15], 1):
                    foods_table.cell(i, 0).text = food.get('user_id', '')[:15]
                    foods_table.cell(i, 1).text = food.get('description', '')[:30]
                    foods_table.cell(i, 2).text = food.get('mealType', '')
                    foods_table.cell(i, 3).text = str(food.get('date', ''))[:10]

            # L∆∞u Word file
            word_buffer = io.BytesIO()
            doc.save(word_buffer)
            word_buffer.seek(0)

            def generate_word():
                yield word_buffer.getvalue()

            return StreamingResponse(
                generate_word(),
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                headers={"Content-Disposition": f"attachment; filename=admin_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"}
            )

        elif format.lower() == "excel":
            if not EXCEL_AVAILABLE:
                # Fallback to CSV if Excel not available
                format = "csv"

        elif format.lower() == "word":
            if not WORD_AVAILABLE:
                # Fallback to CSV if Word not available
                format = "csv"

        # Xu·∫•t CSV (fallback cho t·∫•t c·∫£ formats)
        output = io.StringIO()

        # T·∫°o CSV ƒë∆°n gi·∫£n
        output.write("ADMIN REPORT\n")
        output.write(f"Export Date: {report_data['export_date']}\n")
        output.write(f"Period: {report_data['period']}\n")
        output.write(f"Admin: {report_data['admin']}\n")
        output.write("\n")

        # Summary
        output.write("SUMMARY\n")
        for key, value in report_data['summary'].items():
            output.write(f"{key.replace('_', ' ').title()}: {value}\n")
        output.write("\n")

        # Users data (simplified)
        if users:
            output.write("USERS DATA (Top 10)\n")
            output.write("Email,Name,Created At,Gender\n")

            for user in users:
                try:
                    email = str(user.get('email', 'N/A')).replace(',', ';')[:50]
                    name = str(user.get('name', 'N/A')).replace(',', ';')[:30]
                    created = str(user.get('created_at', 'N/A'))[:19]
                    gender = str(user.get('gender', 'N/A'))[:10]
                    output.write(f'"{email}","{name}","{created}","{gender}"\n')
                except Exception as e:
                    print(f"[EXPORT] Error processing user: {e}")
                    output.write('"Error","Error","Error","Error"\n')
            output.write("\n")
        else:
            output.write("USERS DATA: No users found\n\n")

        # Food records (simplified)
        if food_records:
            output.write("FOOD RECORDS (Top 10)\n")
            output.write("User ID,Description,Meal Type,Date\n")

            for food in food_records:
                try:
                    user_id = str(food.get('user_id', 'N/A'))[:20]
                    desc = str(food.get('description', 'N/A')).replace(',', ';')[:40]
                    meal_type = str(food.get('mealType', 'N/A'))[:15]
                    date = str(food.get('date', 'N/A'))[:19]
                    output.write(f'"{user_id}","{desc}","{meal_type}","{date}"\n')
                except Exception as e:
                    print(f"[EXPORT] Error processing food record: {e}")
                    output.write('"Error","Error","Error","Error"\n')
        else:
            output.write("FOOD RECORDS: No records found\n")

        csv_content = output.getvalue()
        print(f"[EXPORT] CSV content generated, length: {len(csv_content)} characters")

        def generate_csv():
            yield csv_content

        print(f"[EXPORT] Returning StreamingResponse for CSV")
        return StreamingResponse(
            generate_csv(),
            media_type="text/csv",
            headers={"Content-Disposition": f"attachment; filename=admin_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv"}
        )

    except Exception as e:
        print(f"Error exporting report: {str(e)}")
        raise HTTPException(status_code=500, detail=f"L·ªói xu·∫•t b√°o c√°o: {str(e)}")
